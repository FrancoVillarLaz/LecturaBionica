import os
from docx import Document
import pyphen
from datetime import datetime

def format_text(text):
    dic = pyphen.Pyphen(lang='es') 
    formatted_text = []
    for sentence in text:
        formatted_sentence = []
        for word in sentence.split():
            syllables = dic.inserted(word).split("-") 
            first_syllable = syllables[0]
            rest_of_word = ''.join(syllables[1:])
            formatted_sentence.append((first_syllable, rest_of_word))
        formatted_text.append(formatted_sentence)
    return formatted_text


print("Por favor, ingresa el texto plano. Presiona Enter después de cada párrafo. Para finalizar, presiona Enter dos veces:")
user_input = []
while True:
    line = input()
    if not line:
        if len(user_input[-1]) == 0: 
            break
        user_input.append("") 
    else:
        user_input.append(line)


formatted_text = format_text(user_input)


timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
doc_name = f"Lectura_bionica_{timestamp}.docx"

doc = Document()

for sentence in formatted_text:
    p = doc.add_paragraph()
    for part in sentence:
        p.add_run(part[0]).bold = True
        p.add_run(part[1] + " ")


current_directory = os.getcwd()
doc_path = os.path.join(current_directory, doc_name)
doc.save(doc_path)

print("Documento creado en:", doc_path)
